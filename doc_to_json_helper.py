import os
import json
import google.generativeai as genai
from docx import Document


api_key = "AIzaSyDghlziYEo3ZRFO2GwfwA2L0QIUiJ3kj6k"
genai.configure(api_key=api_key)


def extract_text_from_docx(file_path):
    """
    Reads a .docx file and returns the full text content.
    """
    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty lines
                full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_data))

        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading file: {e}")
        return None


def analyze_document_with_gemini(document_text):
    """
    Sends the text to Gemini and requests a specific JSON format.
    """

    model = genai.GenerativeModel("gemini-2.5-flash")

    # The prompt includes the target structure
    prompt = f"""
    You are an expert educational content analyzer.
    Analyze the provided document text and extract the information into a structured JSON format.

    The JSON output must strictly adhere to this schema:
    dont add the ```json markdown text label as ill assume the content is valid json in python
    {{
      "title": "Module Title",
      "description": "Short summary of the document",
      "status": "Draft or Published",
      "objectives": ["List of learning objectives"],
      "sections": [
        {{
          "section_title": "Title of the section",
          "content_type": "Text",
          "body_text": "Summary of the section content, key points, and purpose."
        }}
      ],
      "questions": [
        {{
          "question_text": "Generate a relevant multiple choice question based on the text",
          "difficulty_level": "Easy/Medium/Hard",
          "options": [
            {{ "option_text": "Option A", "is_correct": false }},
            {{ "option_text": "Option B (Correct Answer)", "is_correct": true }}
          ]
        }}
      ]
    }}

    Directives:
    1. Extract 5-8 distinct sections.
    2. Generate 3-5 quiz questions based on the content.
    3. Ensure the JSON is valid.

    Here is the document text to analyze:
    {document_text}
    """

    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error communicating with Gemini: {e}")
        return None
